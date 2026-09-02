#!/usr/bin/env python3
"""
CatchMeUp — recording -> mp3 -> WhisperKit transcript -> Claude recap -> .docx

Modes:
  meeting  — standup / Zoom / client call (decisions, action items)
  lecture  — class / seminar (concepts, terms, study checklist)

  ./catchup meeting <recording>
  ./catchup lecture <recording>
"""
import argparse
import json
import os
import shutil
import subprocess
import sys
import time
import traceback
from datetime import datetime, timedelta
from pathlib import Path

PROJECT_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = PROJECT_DIR / "output"
PROCESSED_DIR = PROJECT_DIR / "processed"
LOGS_DIR = PROJECT_DIR / "logs"
ENV_FILE = PROJECT_DIR / ".env"

CLAUDE_MODEL = "claude-haiku-4-5-20251001"
MAX_RETRIES = 5
MODES = ("meeting", "lecture")

LECTURE_NAME_HINTS = (
    "lecture", "class", "lesson", "seminar", "tutorial", "recitation",
    "course", "week", "midterm", "final", "exam", "homework", "lab",
    "professor", "prof", "ta-", "canvas", "panopto", "kaltura",
)
MEETING_NAME_HINTS = (
    "standup", "stand-up", "stand_up", "zoom", "meet", "huddle", "sync",
    "1-1", "1on1", "1_1", "allhands", "all-hands", "retro", "sprint",
    "kickoff", "interview", "client", "staff", "weekly",
)

MEETING_SCHEMA = """
Return ONLY valid JSON (no markdown fences, no commentary) matching this shape:
{
  "title": "short descriptive meeting title",
  "tldr": ["bullet 1", "bullet 2", "..."],
  "action_items": ["who does what, with deadline if mentioned"],
  "bookmarks": [
    {"timestamp": "HH:MM:SS", "heading": "short label", "insight": "why this moment matters, explained in plain terms"}
  ],
  "detailed_notes": [
    {"heading": "topic heading", "content": "in-depth explanation / notes for this topic, several sentences"}
  ]
}
Include 5-15 bookmarks for the important/decision/action-item moments, spread across the whole meeting.
Include as many detailed_notes sections as needed to cover the meeting thoroughly.
List every follow-up and owner you can hear in action_items.
"""

LECTURE_SCHEMA = """
Return ONLY valid JSON (no markdown fences, no commentary) matching this shape:
{
  "title": "short lecture title (course + topic if you can tell)",
  "tldr": ["the 5-10 things a student who missed class must know"],
  "bookmarks": [
    {"timestamp": "HH:MM:SS", "heading": "concept or example name", "insight": "why this moment matters for understanding or the exam"}
  ],
  "detailed_notes": [
    {"heading": "topic heading", "content": "teach this topic clearly in several sentences, as if the student was absent"}
  ],
  "terms": [
    {"term": "vocab / formula / name", "definition": "plain-language definition"}
  ],
  "study": ["exam-style prompt or thing to memorize / practice"]
}
Include 5-15 bookmarks for definitions, worked examples, and "this will be on the exam" moments.
Cover the lecture thoroughly in detailed_notes. Prefer teaching over quoting.
"""


def log(msg):
    LOGS_DIR.mkdir(exist_ok=True)
    line = f"[{datetime.now().isoformat(timespec='seconds')}] {msg}"
    print(line, flush=True)
    with open(LOGS_DIR / "pipeline.log", "a") as f:
        f.write(line + "\n")


def load_env():
    if ENV_FILE.exists():
        for line in ENV_FILE.read_text().splitlines():
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            os.environ.setdefault(key.strip(), value.strip())


def which_bin(name, env_var, fallbacks):
    from_env = os.environ.get(env_var)
    if from_env and Path(from_env).exists():
        return from_env
    found = shutil.which(name)
    if found:
        return found
    for candidate in fallbacks:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError(
        f"{name} not found. Install it with Homebrew:\n"
        f"  brew install {name}\n"
        f"Then re-run: ./catchup doctor"
    )


def ffmpeg_bin():
    return which_bin(
        "ffmpeg",
        "FFMPEG",
        ["/opt/homebrew/bin/ffmpeg", "/usr/local/bin/ffmpeg"],
    )


def whisperkit_bin():
    return which_bin(
        "whisperkit-cli",
        "WHISPERKIT_CLI",
        ["/opt/homebrew/bin/whisperkit-cli", "/usr/local/bin/whisperkit-cli"],
    )


def notify(title, message):
    if sys.platform != "darwin":
        return
    safe_title = title.replace('"', "'")
    safe_message = message.replace('"', "'")
    subprocess.run(
        ["osascript", "-e", f'display notification "{safe_message}" with title "{safe_title}"'],
        check=False,
    )


def run(cmd, **kwargs):
    log(f"$ {' '.join(str(c) for c in cmd)}")
    subprocess.run(cmd, check=True, **kwargs)


def to_mp3(source: Path) -> Path:
    mp3_path = source.with_suffix(".mp3")
    if mp3_path.exists():
        return mp3_path
    run([
        ffmpeg_bin(), "-y", "-i", str(source),
        "-vn", "-acodec", "libmp3lame", "-ab", "192k", str(mp3_path),
    ])
    return mp3_path


def transcribe(mp3_path: Path) -> dict:
    json_path = mp3_path.with_suffix(".json")
    if not json_path.exists():
        run([
            whisperkit_bin(), "transcribe",
            "--audio-path", str(mp3_path),
            "--language", "en",
            "--report",
            "--report-path", str(mp3_path.parent),
        ])
    if not json_path.exists():
        raise FileNotFoundError(
            f"WhisperKit did not write {json_path.name}. "
            "Run `./catchup doctor` and try again."
        )
    return json.loads(json_path.read_text())


def format_timestamp(seconds: float) -> str:
    return str(timedelta(seconds=int(seconds)))


def transcript_with_timestamps(whisper_json: dict) -> str:
    lines = []
    segments = whisper_json.get("segments") or whisper_json.get("transcription", {}).get("segments") or []
    for seg in segments:
        start = seg.get("start", seg.get("startTime", 0))
        text = (seg.get("text") or "").strip()
        if text:
            lines.append(f"[{format_timestamp(float(start))}] {text}")
    if not lines:
        text = (whisper_json.get("text") or "").strip()
        if text:
            return text
        raise ValueError("Transcript JSON had no segments or text")
    return "\n".join(lines)


def guess_mode(path: Path) -> str:
    name = path.name.lower()
    if any(h in name for h in LECTURE_NAME_HINTS):
        return "lecture"
    if any(h in name for h in MEETING_NAME_HINTS):
        return "meeting"
    env_mode = (os.environ.get("CATCHMEUP_MODE") or "").strip().lower()
    if env_mode in MODES:
        return env_mode
    return "meeting"


def call_claude(transcript_text: str, mode: str) -> dict:
    import anthropic

    if mode == "lecture":
        role = (
            "You are writing notes for a student who missed this recorded lecture. "
            "Teach the material: definitions, examples, formulas, and what is likely to show up on a quiz or exam. "
            "Do not write as if this were a business meeting."
        )
        schema = LECTURE_SCHEMA
    else:
        role = (
            "You are analyzing a transcript of a work meeting (standup, Zoom, client call, 1:1). "
            "Write notes for someone who did not attend: decisions, owners, deadlines, and follow-ups."
        )
        schema = MEETING_SCHEMA

    client = anthropic.Anthropic()
    prompt = f"{role}\n\n{schema}\n\nTranscript (timestamped):\n{transcript_text}"

    delay = 5
    last_err = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=8000,
                messages=[{"role": "user", "content": prompt}],
            )
            raw = resp.content[0].text.strip()
            if raw.startswith("```"):
                raw = raw.strip("`")
                raw = raw.split("\n", 1)[1] if "\n" in raw else raw
            return json.loads(raw)
        except anthropic.RateLimitError as e:
            last_err = e
            log(f"Rate limited (attempt {attempt}/{MAX_RETRIES}), backing off {delay}s")
            time.sleep(delay)
            delay = min(delay * 2, 60)
        except anthropic.APIStatusError as e:
            last_err = e
            if e.status_code >= 500:
                log(f"Claude API {e.status_code} (attempt {attempt}/{MAX_RETRIES}), retrying in {delay}s")
                time.sleep(delay)
                delay = min(delay * 2, 60)
            else:
                raise
    raise RuntimeError(f"Claude API call failed after {MAX_RETRIES} attempts: {last_err}")


def _bullets(doc, items):
    for bullet in items or []:
        doc.add_paragraph(str(bullet), style="List Bullet")


def build_docx(analysis: dict, source_name: str, recorded_at: str, mode: str) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(analysis.get("title", source_name), level=0)
    kind = "Lecture" if mode == "lecture" else "Meeting"
    doc.add_paragraph(f"CatchMeUp {kind.lower()} recap")
    doc.add_paragraph(f"Source recording: {source_name}")
    doc.add_paragraph(f"{kind} date: {recorded_at}")

    tldr_heading = "What you missed" if mode == "lecture" else "TL;DR"
    doc.add_heading(tldr_heading, level=1)
    _bullets(doc, analysis.get("tldr", []))

    if mode == "meeting":
        items = analysis.get("action_items") or []
        if items:
            doc.add_heading("Action items & follow-ups", level=1)
            _bullets(doc, items)

    bookmark_heading = "Key moments" if mode == "lecture" else "Key bookmarks & insights"
    doc.add_heading(bookmark_heading, level=1)
    for bm in analysis.get("bookmarks", []):
        p = doc.add_paragraph()
        run_ts = p.add_run(f"[{bm.get('timestamp', '?')}] ")
        run_ts.bold = True
        run_head = p.add_run(bm.get("heading", ""))
        run_head.bold = True
        run_head.font.size = Pt(12)
        doc.add_paragraph(bm.get("insight", ""))

    notes_heading = "Lecture notes" if mode == "lecture" else "Detailed notes"
    doc.add_heading(notes_heading, level=1)
    for section in analysis.get("detailed_notes", []):
        doc.add_heading(section.get("heading", ""), level=2)
        doc.add_paragraph(section.get("content", ""))

    if mode == "lecture":
        terms = analysis.get("terms") or []
        if terms:
            doc.add_heading("Terms & definitions", level=1)
            for item in terms:
                if isinstance(item, dict):
                    p = doc.add_paragraph()
                    name = p.add_run(f"{item.get('term', '')}: ")
                    name.bold = True
                    p.add_run(item.get("definition", ""))
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        study = analysis.get("study") or []
        if study:
            doc.add_heading("Study / exam checklist", level=1)
            _bullets(doc, study)

    OUTPUT_DIR.mkdir(exist_ok=True)
    suffix = "lecture" if mode == "lecture" else "meeting"
    out_path = OUTPUT_DIR / f"{Path(source_name).stem}_{suffix}_notes.docx"
    doc.save(out_path)
    return out_path


def archive(*paths: Path, dest_dir: Path):
    dest_dir.mkdir(parents=True, exist_ok=True)
    for p in paths:
        if p.exists():
            p.rename(dest_dir / p.name)


def parse_args(argv):
    parser = argparse.ArgumentParser(
        prog="catchup",
        description="Turn a meeting or lecture recording into Word notes.",
    )
    parser.add_argument("recording", help="Path to the video/audio file")
    parser.add_argument(
        "--mode",
        choices=MODES,
        help="meeting (work) or lecture (class). Guessed from the filename if omitted.",
    )
    return parser.parse_args(argv)


def main(argv=None):
    args = parse_args(argv if argv is not None else sys.argv[1:])
    load_env()
    source = Path(args.recording).resolve()
    if not source.exists():
        log(f"ERROR: source file not found: {source}")
        sys.exit(1)

    mode = args.mode or guess_mode(source)
    if mode not in MODES:
        mode = "meeting"

    if not os.environ.get("ANTHROPIC_API_KEY"):
        log("ERROR: ANTHROPIC_API_KEY not set — run ./catchup config")
        notify("CatchMeUp", "Missing ANTHROPIC_API_KEY - see logs/pipeline.log")
        sys.exit(1)

    try:
        log(f"Processing {source.name} as {mode}")
        mp3_path = to_mp3(source)
        whisper_json = transcribe(mp3_path)
        transcript_text = transcript_with_timestamps(whisper_json)

        log(f"Calling Claude for {mode} recap...")
        analysis = call_claude(transcript_text, mode)

        recorded_at = datetime.fromtimestamp(source.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
        docx_path = build_docx(analysis, source.name, recorded_at, mode)
        log(f"Wrote {docx_path}")

        archive_dir = PROCESSED_DIR / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{source.stem}"
        json_path = mp3_path.with_suffix(".json")
        srt_path = mp3_path.with_suffix(".srt")
        archive(source, mp3_path, json_path, srt_path, dest_dir=archive_dir)
        log(f"Archived source files to {archive_dir}")
        notify("CatchMeUp", f"Done: {docx_path.name} is ready in output/")

    except Exception:
        log("ERROR: pipeline failed\n" + traceback.format_exc())
        notify("CatchMeUp FAILED", f"{source.name} - see logs/pipeline.log")
        sys.exit(1)


if __name__ == "__main__":
    main()
