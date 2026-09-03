#!/usr/bin/env python3
"""
CatchMeUp — recording -> mp3 -> WhisperKit transcript -> LLM recap -> .docx

Modes:
  meeting  — standup / Zoom / client call (decisions, action items)
  lecture  — class / seminar (concepts, terms, study checklist)

  ./catchup meeting <recording>
  ./catchup lecture <recording>
"""
import argparse
import json
import os
import re
import shutil
import subprocess
import sys
import traceback
from datetime import datetime, timedelta
from pathlib import Path

from . import brains as brains_mod

MODES = ("meeting", "lecture")

LECTURE_NAME_HINTS = (
    "lecture", "class", "lesson", "seminar", "tutorial", "recitation",
    "course", "week", "midterm", "final", "exam", "homework", "lab",
    "professor", "prof", "ta-", "canvas", "panopto", "kaltura",
)
MEETING_NAME_HINTS = (
    "standup", "stand-up", "stand_up", "zoom", "meet", "huddle", "sync",
    "1-1", "1on1", "1_1", "allhands", "all-hands", "retro", "sprint",
    "kickoff", "interview", "client", "staff", "weekly",
)

MEETING_SCHEMA = """
Return ONLY valid JSON (no markdown fences, no commentary) matching this shape:
{
  "title": "short descriptive meeting title",
  "tldr": ["bullet 1", "bullet 2", "..."],
  "action_items": ["Speaker 1 / Jordan: who does what, with deadline if mentioned"],
  "speakers": [
    {"label": "Speaker 1", "name": "Jordan or unknown", "said": "their role in this meeting in one line"}
  ],
  "bookmarks": [
    {"timestamp": "HH:MM:SS", "heading": "short label", "insight": "why this moment matters, explained in plain terms"}
  ],
  "detailed_notes": [
    {"heading": "topic heading", "content": "in-depth explanation / notes for this topic, several sentences"}
  ]
}
Include 5-15 bookmarks for the important/decision/action-item moments, spread across the whole meeting.
Include as many detailed_notes sections as needed to cover the meeting thoroughly.
List every follow-up and owner you can hear in action_items. Prefer the speaker label (Speaker 1) plus any name they used.
If the transcript has speaker labels, fill speakers[] — do not invent names you cannot hear.
"""

LECTURE_SCHEMA = """
Return ONLY valid JSON (no markdown fences, no commentary) matching this shape:
{
  "title": "short lecture title (course + topic if you can tell)",
  "tldr": ["the 5-10 things a student who missed class must know"],
  "bookmarks": [
    {"timestamp": "HH:MM:SS", "heading": "concept or example name", "insight": "why this moment matters for understanding or the exam"}
  ],
  "detailed_notes": [
    {"heading": "topic heading", "content": "teach this topic clearly in several sentences, as if the student was absent"}
  ],
  "terms": [
    {"term": "vocab / formula / name", "definition": "plain-language definition"}
  ],
  "study": ["exam-style prompt or thing to memorize / practice"]
}
Include 5-15 bookmarks for definitions, worked examples, and "this will be on the exam" moments.
Cover the lecture thoroughly in detailed_notes. Prefer teaching over quoting.
"""


def output_dir() -> Path:
    return brains_mod.output_root()


def processed_dir() -> Path:
    return brains_mod.processed_root()


def logs_dir() -> Path:
    return brains_mod.logs_root()


def log(msg):
    logs_dir().mkdir(parents=True, exist_ok=True)
    line = f"[{datetime.now().isoformat(timespec='seconds')}] {msg}"
    print(line, flush=True)
    with open(logs_dir() / "pipeline.log", "a") as f:
        f.write(line + "\n")


def load_env():
    brains_mod.load_env()


def which_bin(name, env_var, fallbacks):
    from_env = os.environ.get(env_var)
    if from_env and Path(from_env).exists():
        return from_env
    found = shutil.which(name)
    if found:
        return found
    for candidate in fallbacks:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError(
        f"{name} not found. Install it with Homebrew:\n"
        f"  brew install {name}\n"
        f"Then re-run: ./catchup doctor"
    )


def ffmpeg_bin():
    return which_bin(
        "ffmpeg",
        "FFMPEG",
        ["/opt/homebrew/bin/ffmpeg", "/usr/local/bin/ffmpeg"],
    )


def whisperkit_bin():
    return which_bin(
        "whisperkit-cli",
        "WHISPERKIT_CLI",
        ["/opt/homebrew/bin/whisperkit-cli", "/usr/local/bin/whisperkit-cli"],
    )


def notify(title, message):
    if sys.platform != "darwin":
        return
    safe_title = title.replace('"', "'")
    safe_message = message.replace('"', "'")
    subprocess.run(
        ["osascript", "-e", f'display notification "{safe_message}" with title "{safe_title}"'],
        check=False,
    )


def run(cmd, **kwargs):
    log(f"$ {' '.join(str(c) for c in cmd)}")
    subprocess.run(cmd, check=True, **kwargs)


def to_mp3(source: Path) -> Path:
    mp3_path = source.with_suffix(".mp3")
    if mp3_path.exists():
        return mp3_path
    run([
        ffmpeg_bin(), "-y", "-i", str(source),
        "-vn", "-acodec", "libmp3lame", "-ab", "192k", str(mp3_path),
    ])
    return mp3_path


def want_diarize(mode: str) -> bool:
    raw = (os.environ.get("CATCHMEUP_DIARIZE") or "").strip().lower()
    if raw in {"0", "false", "no", "off"}:
        return False
    if raw in {"1", "true", "yes", "on", "always"}:
        return True
    return mode == "meeting"


def pretty_speaker(raw: str) -> str:
    """Turn SPEAKER_00 / spk_1 / 0 into a stable 'Speaker N' label."""
    s = (raw or "").strip()
    if not s:
        return ""
    named = re.match(r"(?i)^speaker\s+(\d+)$", s)
    if named:
        return f"Speaker {int(named.group(1))}"
    digits = re.search(r"(\d+)$", s)
    if digits and re.search(r"(?i)(speaker|spk)", s):
        return f"Speaker {int(digits.group(1)) + 1}"
    if re.match(r"^\d+$", s):
        n = int(s)
        return f"Speaker {n + 1 if n == 0 else n}"
    return s


def segment_speaker(seg: dict) -> str:
    if not isinstance(seg, dict):
        return ""
    for key in ("speaker", "speakerLabel", "speaker_label", "speakerId", "speaker_id"):
        val = seg.get(key)
        if val is not None and str(val).strip() != "":
            return pretty_speaker(str(val))
    words = seg.get("words") or seg.get("tokens") or []
    votes: dict[str, int] = {}
    for w in words:
        if not isinstance(w, dict):
            continue
        sp = w.get("speaker") or w.get("speakerLabel")
        if sp:
            label = pretty_speaker(str(sp))
            votes[label] = votes.get(label, 0) + 1
    if votes:
        return max(votes, key=votes.get)
    return ""


def parse_rttm(path: Path) -> list[tuple[float, float, str]]:
    turns = []
    if not path.is_file():
        return turns
    for line in path.read_text().splitlines():
        parts = line.split()
        if len(parts) < 8 or parts[0] != "SPEAKER":
            continue
        try:
            start = float(parts[3])
            dur = float(parts[4])
        except ValueError:
            continue
        turns.append((start, start + dur, pretty_speaker(parts[7])))
    return turns


def apply_rttm(whisper_json: dict, rttm_path: Path) -> dict:
    turns = parse_rttm(rttm_path)
    if not turns:
        return whisper_json
    segments = whisper_json.get("segments") or whisper_json.get("transcription", {}).get("segments") or []
    for seg in segments:
        if segment_speaker(seg):
            continue
        start = float(seg.get("start", seg.get("startTime", 0)) or 0)
        best = None
        best_overlap = 0.0
        end = float(seg.get("end", seg.get("endTime", start + 1)) or (start + 1))
        for t0, t1, spk in turns:
            overlap = min(end, t1) - max(start, t0)
            if overlap > best_overlap:
                best_overlap = overlap
                best = spk
        if best:
            seg["speaker"] = best
    return whisper_json


def transcribe(mp3_path: Path, diarize: bool = False) -> dict:
    json_path = mp3_path.with_suffix(".json")
    rttm_path = mp3_path.with_suffix(".rttm")
    if not json_path.exists():
        cmd = [
            whisperkit_bin(), "transcribe",
            "--audio-path", str(mp3_path),
            "--language", "en",
            "--report",
            "--report-path", str(mp3_path.parent),
        ]
        if diarize:
            cmd.append("--diarization")
            log("Diarizing speakers (meeting mode)")
        try:
            run(cmd)
        except subprocess.CalledProcessError:
            if diarize:
                log("Diarization failed — transcribing without speaker labels")
                run([
                    whisperkit_bin(), "transcribe",
                    "--audio-path", str(mp3_path),
                    "--language", "en",
                    "--report",
                    "--report-path", str(mp3_path.parent),
                ])
            else:
                raise
    if not json_path.exists():
        raise FileNotFoundError(
            f"WhisperKit did not write {json_path.name}. "
            "Run `./catchup doctor` and try again."
        )
    data = json.loads(json_path.read_text())
    segments = data.get("segments") or data.get("transcription", {}).get("segments") or []
    has_speakers = any(segment_speaker(s) for s in segments)
    if diarize and not has_speakers:
        try:
            run([
                whisperkit_bin(), "diarize",
                "--audio-path", str(mp3_path),
                "--rttm-path", str(rttm_path),
            ])
            data = apply_rttm(data, rttm_path)
            json_path.write_text(json.dumps(data, indent=2) + "\n")
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            log(f"Speaker RTTM skipped: {e}")
    return data


def format_timestamp(seconds: float) -> str:
    return str(timedelta(seconds=int(seconds)))


def transcript_with_timestamps(whisper_json: dict) -> str:
    lines = []
    segments = whisper_json.get("segments") or whisper_json.get("transcription", {}).get("segments") or []
    for seg in segments:
        start = seg.get("start", seg.get("startTime", 0))
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        speaker = segment_speaker(seg)
        stamp = format_timestamp(float(start))
        if speaker:
            lines.append(f"[{stamp}] {speaker}: {text}")
        else:
            lines.append(f"[{stamp}] {text}")
    if not lines:
        text = (whisper_json.get("text") or "").strip()
        if text:
            return text
        raise ValueError("Transcript JSON had no segments or text")
    return "\n".join(lines)


def guess_mode(path: Path) -> str:
    name = path.name.lower()
    if any(h in name for h in LECTURE_NAME_HINTS):
        return "lecture"
    if any(h in name for h in MEETING_NAME_HINTS):
        return "meeting"
    env_mode = (os.environ.get("CATCHMEUP_MODE") or "").strip().lower()
    if env_mode in MODES:
        return env_mode
    return "meeting"


def call_llm(transcript_text: str, mode: str, speaker_hint: str = "") -> dict:
    from .providers import complete_json

    if mode == "lecture":
        role = (
            "You are writing notes for a student who missed this recorded lecture. "
            "Teach the material: definitions, examples, formulas, and what is likely to show up on a quiz or exam. "
            "Do not write as if this were a business meeting."
        )
        schema = LECTURE_SCHEMA
    else:
        role = (
            "You are analyzing a transcript of a work meeting (standup, Zoom, client call, 1:1). "
            "Write notes for someone who did not attend: decisions, owners, deadlines, and follow-ups. "
            "Lines may be labeled Speaker 1, Speaker 2, … — use those labels in action items and speakers[]."
        )
        schema = MEETING_SCHEMA
        if speaker_hint:
            role += speaker_hint

    prompt = f"{role}\n\n{schema}\n\nTranscript (timestamped):\n{transcript_text}"
    return complete_json(prompt, log=log)


def _bullets(doc, items):
    for bullet in items or []:
        doc.add_paragraph(str(bullet), style="List Bullet")


def build_docx(analysis: dict, source_name: str, recorded_at: str, mode: str) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(analysis.get("title", source_name), level=0)
    kind = "Lecture" if mode == "lecture" else "Meeting"
    doc.add_paragraph(f"CatchMeUp {kind.lower()} recap")
    doc.add_paragraph(f"Source recording: {source_name}")
    doc.add_paragraph(f"{kind} date: {recorded_at}")

    tldr_heading = "What you missed" if mode == "lecture" else "TL;DR"
    doc.add_heading(tldr_heading, level=1)
    _bullets(doc, analysis.get("tldr", []))

    if mode == "meeting":
        speakers = analysis.get("speakers") or []
        if speakers:
            doc.add_heading("Who spoke", level=1)
            for sp in speakers:
                if isinstance(sp, dict):
                    label = sp.get("name") or sp.get("label") or "Speaker"
                    said = (sp.get("said") or "").strip()
                    p = doc.add_paragraph()
                    name = p.add_run(f"{label}: ")
                    name.bold = True
                    p.add_run(said)
                else:
                    doc.add_paragraph(str(sp), style="List Bullet")
        items = analysis.get("action_items") or []
        if items:
            doc.add_heading("Action items & follow-ups", level=1)
            _bullets(doc, items)

    bookmark_heading = "Key moments" if mode == "lecture" else "Key bookmarks & insights"
    doc.add_heading(bookmark_heading, level=1)
    for bm in analysis.get("bookmarks", []):
        p = doc.add_paragraph()
        run_ts = p.add_run(f"[{bm.get('timestamp', '?')}] ")
        run_ts.bold = True
        run_head = p.add_run(bm.get("heading", ""))
        run_head.bold = True
        run_head.font.size = Pt(12)
        doc.add_paragraph(bm.get("insight", ""))

    notes_heading = "Lecture notes" if mode == "lecture" else "Detailed notes"
    doc.add_heading(notes_heading, level=1)
    for section in analysis.get("detailed_notes", []):
        doc.add_heading(section.get("heading", ""), level=2)
        doc.add_paragraph(section.get("content", ""))

    if mode == "lecture":
        terms = analysis.get("terms") or []
        if terms:
            doc.add_heading("Terms & definitions", level=1)
            for item in terms:
                if isinstance(item, dict):
                    p = doc.add_paragraph()
                    name = p.add_run(f"{item.get('term', '')}: ")
                    name.bold = True
                    p.add_run(item.get("definition", ""))
                else:
                    doc.add_paragraph(str(item), style="List Bullet")
        study = analysis.get("study") or []
        if study:
            doc.add_heading("Study / exam checklist", level=1)
            _bullets(doc, study)

    dest = output_dir()
    dest.mkdir(parents=True, exist_ok=True)
    suffix = "lecture" if mode == "lecture" else "meeting"
    out_path = dest / f"{Path(source_name).stem}_{suffix}_notes.docx"
    doc.save(out_path)
    return out_path


def _md_bullets(items) -> list[str]:
    return [f"- {item}" for item in (items or [])]


def _term_md(item) -> str:
    if isinstance(item, dict):
        term = str(item.get("term") or "").strip()
        definition = str(item.get("definition") or "").strip()
        if term:
            return f"- **[[{term}]]:** {definition}"
        return f"- {definition}"
    return f"- {item}"


def render_markdown(analysis: dict, source_name: str, recorded_at: str, mode: str) -> str:
    kind = "Lecture" if mode == "lecture" else "Meeting"
    lines = [
        f"# {analysis.get('title', source_name)}",
        "",
        f"CatchMeUp {kind.lower()} recap  ",
        f"Source: `{source_name}`  ",
        f"{kind} date: {recorded_at}",
        "",
        f"## {'What you missed' if mode == 'lecture' else 'TL;DR'}",
        "",
        *_md_bullets(analysis.get("tldr")),
        "",
    ]
    if mode == "meeting" and analysis.get("speakers"):
        lines += ["## Who spoke", ""]
        for sp in analysis.get("speakers") or []:
            if isinstance(sp, dict):
                label = sp.get("name") or sp.get("label") or "Speaker"
                said = (sp.get("said") or "").strip()
                lines.append(f"- **{label}:** {said}" if said else f"- **{label}**")
            else:
                lines.append(f"- {sp}")
        lines.append("")
    if mode == "meeting" and analysis.get("action_items"):
        lines += ["## Action items & follow-ups", ""] + _md_bullets(analysis.get("action_items")) + [""]
    lines += [f"## {'Key moments' if mode == 'lecture' else 'Key bookmarks & insights'}", ""]
    for bm in analysis.get("bookmarks") or []:
        lines.append(f"**[{bm.get('timestamp', '?')}] {bm.get('heading', '')}**  ")
        lines.append(f"{bm.get('insight', '')}")
        lines.append("")
    lines += [f"## {'Lecture notes' if mode == 'lecture' else 'Detailed notes'}", ""]
    for section in analysis.get("detailed_notes") or []:
        lines.append(f"### {section.get('heading', '')}")
        lines.append("")
        lines.append(section.get("content", ""))
        lines.append("")
    if mode == "lecture":
        if analysis.get("terms"):
            lines += ["## Terms & definitions", ""]
            for item in analysis.get("terms") or []:
                lines.append(_term_md(item))
            lines.append("")
        if analysis.get("study"):
            lines += ["## Study / exam checklist", ""] + _md_bullets(analysis.get("study")) + [""]
    return "\n".join(lines).strip() + "\n"


def build_markdown(analysis: dict, source_name: str, recorded_at: str, mode: str) -> Path:
    dest = output_dir()
    dest.mkdir(parents=True, exist_ok=True)
    suffix = "lecture" if mode == "lecture" else "meeting"
    out_path = dest / f"{Path(source_name).stem}_{suffix}_notes.md"
    out_path.write_text(render_markdown(analysis, source_name, recorded_at, mode))
    return out_path


def persist_recap(
    analysis: dict,
    source: Path,
    recorded_at: str,
    mode: str,
    transcript_text: str,
    provider: str,
    brain_slug: str | None = None,
    md_path: Path | None = None,
    docx_path: Path | None = None,
) -> tuple[dict, Path]:
    """Write notes + catchmeup.json (+ cortex ingest). Used by the pipeline and tests."""
    if md_path is None:
        md_path = build_markdown(analysis, source.name, recorded_at, mode)
    if docx_path is None:
        try:
            docx_path = build_docx(analysis, source.name, recorded_at, mode)
        except ImportError:
            docx_path = None

    if brain_slug:
        brains_mod.load_brain(brain_slug)
        archive_dir = brains_mod.unique_stamp_dir(brains_mod.recaps_dir(brain_slug), source.stem)
        notes = brains_mod.notes_dir(brain_slug)
        notes.mkdir(parents=True, exist_ok=True)
        if md_path and md_path.exists():
            (notes / md_path.name).write_text(md_path.read_text())
    else:
        archive_dir = brains_mod.unique_stamp_dir(processed_dir(), source.stem)

    archive_dir.mkdir(parents=True, exist_ok=True)
    (archive_dir / "transcript.txt").write_text(transcript_text)
    record = {
        "version": 1,
        "mode": mode,
        "brain": brain_slug,
        "title": analysis.get("title") or source.stem,
        "source": source.name,
        "source_path": str(source.resolve()),
        "recorded_at": recorded_at,
        "processed_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "provider": provider,
        "audio": source.with_suffix(".mp3").name,
        "markdown": str(md_path.name) if md_path else "",
        "docx": str(docx_path.name) if docx_path else "",
        "analysis": analysis,
    }
    (archive_dir / "catchmeup.json").write_text(json.dumps(record, indent=2) + "\n")
    if brain_slug:
        from . import cortex as cortex_mod

        cortex_mod.ingest_recap(brain_slug, record)
    record["_dir"] = str(archive_dir)
    return record, archive_dir


def archive_pipeline_outputs(source: Path, mp3_path: Path, dest_dir: Path, keep: bool) -> None:
    """Move whisper sidecars into the recap folder. Optionally leave the original media."""
    dest_dir.mkdir(parents=True, exist_ok=True)
    for p in (
        mp3_path,
        mp3_path.with_suffix(".json"),
        mp3_path.with_suffix(".srt"),
        mp3_path.with_suffix(".rttm"),
    ):
        if p.exists():
            dest = dest_dir / p.name
            if p.resolve() != dest.resolve():
                p.rename(dest)
    if keep:
        return
    dest = dest_dir / source.name
    if source.exists() and source.resolve() != dest.resolve():
        source.rename(dest)


def ingest_limit() -> int | None:
    raw = (os.environ.get("CATCHMEUP_INGEST_LIMIT") or "").strip()
    if not raw:
        return None
    try:
        n = int(raw)
    except ValueError:
        return None
    return n if n > 0 else None


def process_recording(source: Path, mode: str | None, brain_slug: str | None) -> None:
    from .providers import active_provider, resolve_api_key
    from . import viz

    source = source.resolve()
    mode = mode or guess_mode(source)
    if mode not in MODES:
        mode = "meeting"
    provider = active_provider()
    if not resolve_api_key(provider) and provider != "ollama":
        raise RuntimeError(f"no API key for {provider} — run ./catchup config {provider}")

    if brain_slug and source.name in brains_mod.ingested_sources(brain_slug):
        log(f"Skip (already in {brain_slug}): {source.name}")
        print(f"Already filed: {source.name}", flush=True)
        return

    print(viz.pipeline_track("audio", f"{source.name}  ·  {mode}"), flush=True)
    log(f"Processing {source.name} as {mode}")
    mp3_path = to_mp3(source)
    print(viz.pipeline_track("whisper", "on-device transcript"), flush=True)
    whisper_json = transcribe(mp3_path, diarize=want_diarize(mode))
    transcript_text = transcript_with_timestamps(whisper_json)

    print(viz.pipeline_track("llm", f"{provider} recap"), flush=True)
    log(f"Calling {provider} for {mode} recap...")
    hint = brains_mod.speaker_prompt_hint(brain_slug) if brain_slug else ""
    analysis = call_llm(transcript_text, mode, speaker_hint=hint)
    if brain_slug:
        analysis = brains_mod.apply_speaker_map(brain_slug, analysis)

    print(viz.pipeline_track("notes"), flush=True)
    recorded_at = datetime.fromtimestamp(source.stat().st_mtime).strftime("%Y-%m-%d %H:%M")
    md_path = build_markdown(analysis, source.name, recorded_at, mode)
    docx_path = build_docx(analysis, source.name, recorded_at, mode)
    log(f"Wrote {docx_path}")
    log(f"Wrote {md_path}")

    record, archive_dir = persist_recap(
        analysis,
        source,
        recorded_at,
        mode,
        transcript_text,
        provider,
        brain_slug=brain_slug,
        md_path=md_path,
        docx_path=docx_path,
    )
    if brain_slug:
        log(f"Cortex updated for brain {brain_slug}")
    keep = brains_mod.keep_source(source, brain_slug)
    archive_pipeline_outputs(source, mp3_path, archive_dir, keep=keep)
    log(f"Archived artifacts to {archive_dir}" + (" (original kept)" if keep else ""))
    print(viz.pipeline_track("done"), flush=True)
    print(viz.recap_card(analysis, mode, md_path, docx_path), flush=True)
    notify("CatchMeUp", f"Done: {docx_path.name} is ready in output/")


def process_folder(folder: Path, mode: str | None, brain_slug: str | None) -> int:
    if not brain_slug:
        log("ERROR: folder ingest needs --brain (./catchup into NAME DIR)")
        print("Folder ingest needs a brain: ./catchup into mit-60001 MIT-6.0001/", flush=True)
        return 1
    brains_mod.load_brain(brain_slug)
    pending = brains_mod.pending_media(brain_slug, folder)
    limit = ingest_limit()
    if limit is not None:
        pending = pending[:limit]
    if not pending:
        print("Nothing new to file (already ingested, or no media in that folder).", flush=True)
        return 0
    failed = 0
    for i, path in enumerate(pending, 1):
        print(f"\n[{i}/{len(pending)}] {path.name}", flush=True)
        try:
            process_recording(path, mode, brain_slug)
        except Exception:
            failed += 1
            log(f"ERROR: {path.name}\n" + traceback.format_exc())
            notify("CatchMeUp FAILED", f"{path.name} - see logs/pipeline.log")
            print(f"Failed: {path.name} — continuing", flush=True)
    if failed:
        log(f"Folder ingest finished with {failed} failure(s) of {len(pending)}")
        return 1
    return 0


def parse_args(argv):
    parser = argparse.ArgumentParser(
        prog="catchup",
        description="Turn a meeting or lecture recording into Word notes.",
    )
    parser.add_argument("recording", help="Path to the video/audio file, or a folder of them")
    parser.add_argument(
        "--mode",
        choices=MODES,
        help="meeting (work) or lecture (class). Guessed from the filename if omitted.",
    )
    parser.add_argument(
        "--brain",
        help="File this recap into a specialist brain folder (./catchup brain new NAME).",
    )
    return parser.parse_args(argv)


def main(argv=None):
    args = parse_args(argv if argv is not None else sys.argv[1:])
    load_env()
    source = Path(args.recording).expanduser()
    if not source.exists():
        log(f"ERROR: source file not found: {source}")
        sys.exit(1)

    brain_slug = (args.brain or "").strip() or None
    if source.is_dir():
        sys.exit(process_folder(source, args.mode, brain_slug))

    from .providers import active_provider, resolve_api_key

    provider = active_provider()
    if not resolve_api_key(provider) and provider != "ollama":
        log(f"ERROR: no API key for {provider} — run ./catchup config {provider}")
        notify("CatchMeUp", "Missing API key - run ./catchup config")
        sys.exit(1)

    try:
        process_recording(source, args.mode, brain_slug)
    except Exception:
        log("ERROR: pipeline failed\n" + traceback.format_exc())
        notify("CatchMeUp FAILED", f"{source.name} - see logs/pipeline.log")
        sys.exit(1)


if __name__ == "__main__":
    main()
